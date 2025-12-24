import os
from pdf2docx import Converter
from docx2pdf import convert as docx_to_pdf
from docx import Document
import pdfplumber
from reportlab.pdfgen import canvas


# ---------------- PDF → DOCX ----------------
def pdf_to_docx(input_file, output_file):
    cv = Converter(input_file)
    cv.convert(output_file, start=0, end=None)
    cv.close()


# ---------------- PDF → TXT ----------------
def pdf_to_txt(input_file, output_file):
    text = ""
    with pdfplumber.open(input_file) as pdf:
        for page in pdf.pages:
            if page.extract_text():
                text += page.extract_text() + "\n"

    with open(output_file, "w", encoding="utf-8") as f:
        f.write(text)


# ---------------- DOCX → TXT ----------------
def docx_to_txt(input_file, output_file):
    doc = Document(input_file)
    with open(output_file, "w", encoding="utf-8") as f:
        for para in doc.paragraphs:
            f.write(para.text + "\n")


# ---------------- TXT → DOCX ----------------
def txt_to_docx(input_file, output_file):
    doc = Document()
    with open(input_file, "r", encoding="utf-8") as f:
        for line in f:
            doc.add_paragraph(line.strip())
    doc.save(output_file)


# ---------------- TXT → PDF ----------------
def txt_to_pdf(input_file, output_file):
    c = canvas.Canvas(output_file)
    y = 800

    with open(input_file, "r", encoding="utf-8") as f:
        for line in f:
            c.drawString(40, y, line.strip())
            y -= 15
            if y < 40:
                c.showPage()
                y = 800

    c.save()


# ---------------- UNIVERSAL CONVERTER ----------------
def convert_file(input_path, output_path):
    input_ext = os.path.splitext(input_path)[1].lower()
    output_ext = os.path.splitext(output_path)[1].lower()

    if input_ext == ".pdf" and output_ext == ".docx":
        pdf_to_docx(input_path, output_path)

    elif input_ext == ".pdf" and output_ext == ".txt":
        pdf_to_txt(input_path, output_path)

    elif input_ext == ".docx" and output_ext == ".pdf":
        docx_to_pdf(input_path, output_path)

    elif input_ext == ".docx" and output_ext == ".txt":
        docx_to_txt(input_path, output_path)

    elif input_ext == ".txt" and output_ext == ".docx":
        txt_to_docx(input_path, output_path)

    elif input_ext == ".txt" and output_ext == ".pdf":
        txt_to_pdf(input_path, output_path)

    else:
        print("❌ Conversion not supported!")
        return

    print(f"✅ Converted '{input_path}' → '{output_path}' successfully.")


# ---------------- MAIN FUNCTION ----------------
def main():
    print("==== Universal File Converter ====")
    input_path = input("Enter input file path: ")
    output_path = input("Enter output file path: ")

    if not os.path.exists(input_path):
        print("❌ Input file does not exist!")
        return

    convert_file(input_path, output_path)


# ---------------- PROGRAM START ----------------
if __name__ == "__main__":
    main()

